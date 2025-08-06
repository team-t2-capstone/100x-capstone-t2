"""
Text Extraction Utilities
Handles extracting text content from various document formats
"""
import os
import re
import logging
from typing import Dict, Optional, Tuple
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document
import markdown

import structlog

logger = structlog.get_logger()

class TextExtractor:
    """Unified text extraction for multiple document formats"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.doc', '.docx', '.md', '.rtf'}
    
    def __init__(self):
        self.extractors = {
            '.pdf': self._extract_from_pdf,
            '.txt': self._extract_from_txt,
            '.docx': self._extract_from_docx,
            '.md': self._extract_from_markdown,
            '.rtf': self._extract_from_rtf,
        }
    
    async def extract_text(self, file_path: str, file_type: str) -> Dict[str, any]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            file_type: File extension (e.g., 'pdf', 'docx')
        
        Returns:
            Dictionary containing:
            - text: Extracted text content
            - metadata: Document metadata (title, pages, etc.)
            - success: Whether extraction succeeded
            - error: Error message if extraction failed
        """
        try:
            # Normalize file type
            file_ext = f".{file_type.lower()}" if not file_type.startswith('.') else file_type.lower()
            
            if file_ext not in self.SUPPORTED_EXTENSIONS:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {file_type}',
                    'text': '',
                    'metadata': {}
                }
            
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'error': f'File not found: {file_path}',
                    'text': '',
                    'metadata': {}
                }
            
            # Get file info
            file_stat = os.stat(file_path)
            base_metadata = {
                'file_size': file_stat.st_size,
                'file_name': os.path.basename(file_path),
                'file_type': file_ext,
            }
            
            logger.info("Extracting text from document", 
                       file_path=file_path, 
                       file_type=file_ext,
                       file_size=file_stat.st_size)
            
            # Extract using appropriate extractor
            extractor = self.extractors.get(file_ext, self._extract_from_txt)
            text, doc_metadata = await extractor(file_path)
            
            # Combine metadata
            metadata = {**base_metadata, **doc_metadata}
            
            # Clean and validate text
            cleaned_text = self._clean_text(text)
            
            if not cleaned_text.strip():
                return {
                    'success': False,
                    'error': 'No text content found in document',
                    'text': '',
                    'metadata': metadata
                }
            
            logger.info("Text extraction successful", 
                       text_length=len(cleaned_text),
                       metadata=metadata)
            
            return {
                'success': True,
                'error': None,
                'text': cleaned_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error("Text extraction failed", error=str(e), file_path=file_path)
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF files"""
        text_content = []
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Get PDF metadata
                metadata.update({
                    'pages': len(pdf_reader.pages),
                    'pdf_metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                })
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning("Failed to extract text from PDF page", 
                                     page_num=page_num + 1, error=str(e))
                        continue
                
                return '\n\n'.join(text_content), metadata
                
        except Exception as e:
            logger.error("PDF text extraction failed", error=str(e))
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_from_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from plain text files"""
        metadata = {}
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                    metadata['encoding'] = encoding
                    metadata['lines'] = len(content.splitlines())
                    
                    return content, metadata
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try binary mode
            with open(file_path, 'rb') as file:
                content = file.read().decode('utf-8', errors='ignore')
            
            metadata['encoding'] = 'binary_fallback'
            metadata['lines'] = len(content.splitlines())
            
            return content, metadata
            
        except Exception as e:
            logger.error("Text file extraction failed", error=str(e))
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    async def _extract_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_texts.append(' | '.join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_texts
            content = '\n\n'.join(all_text)
            
            # Build metadata
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'total_elements': len(all_text)
            }
            
            # Try to get document properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                })
            except Exception:
                pass  # Properties might not be available
            
            return content, metadata
            
        except Exception as e:
            logger.error("DOCX text extraction failed", error=str(e))
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_from_markdown(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Convert markdown to plain text (removes formatting)
            html = markdown.markdown(content)
            
            # Simple HTML tag removal
            plain_text = re.sub(r'<[^>]+>', '', html)
            plain_text = re.sub(r'\n\s*\n', '\n\n', plain_text)
            
            # Build metadata
            metadata = {
                'original_format': 'markdown',
                'lines': len(content.splitlines()),
                'has_headers': bool(re.search(r'^#{1,6}\s', content, re.MULTILINE)),
                'has_links': bool(re.search(r'\[.*?\]\(.*?\)', content)),
                'has_images': bool(re.search(r'!\[.*?\]\(.*?\)', content)),
            }
            
            return plain_text, metadata
            
        except Exception as e:
            logger.error("Markdown text extraction failed", error=str(e))
            raise Exception(f"Failed to extract text from Markdown: {str(e)}")
    
    async def _extract_from_rtf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from RTF files (basic implementation)"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Basic RTF parsing - remove RTF control codes
            # This is a simplified approach; for production, consider using striprtf library
            text = re.sub(r'\\[a-z]+\d*', '', content)  # Remove control words
            text = re.sub(r'[{}]', '', text)  # Remove braces
            text = re.sub(r'\\', '', text)  # Remove remaining backslashes
            text = re.sub(r'\n\s*\n', '\n\n', text)  # Clean up spacing
            
            metadata = {
                'original_format': 'rtf',
                'processed_with': 'basic_regex_parser'
            }
            
            return text.strip(), metadata
            
        except Exception as e:
            logger.error("RTF text extraction failed", error=str(e))
            raise Exception(f"Failed to extract text from RTF: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'\f', '\n', text)  # Form feeds
        text = re.sub(r'\r\n', '\n', text)  # Windows line endings
        text = re.sub(r'\r', '\n', text)  # Mac line endings
        
        # Remove excessive spaces
        text = re.sub(r'  +', ' ', text)
        
        # Clean up line breaks
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def get_text_statistics(self, text: str) -> Dict[str, int]:
        """Get basic statistics about extracted text"""
        if not text:
            return {
                'characters': 0,
                'words': 0,
                'lines': 0,
                'paragraphs': 0
            }
        
        lines = text.split('\n')
        paragraphs = [line for line in lines if line.strip()]
        words = text.split()
        
        return {
            'characters': len(text),
            'words': len(words),
            'lines': len(lines),
            'paragraphs': len(paragraphs)
        }

# Global text extractor instance
text_extractor = TextExtractor()

# Convenience functions
async def extract_text_from_file(file_path: str, file_type: str) -> Dict[str, any]:
    """Extract text from a file using the global text extractor"""
    return await text_extractor.extract_text(file_path, file_type)

def validate_document_format(file_path: str, file_type: str) -> bool:
    """Check if a document format is supported"""
    file_ext = f".{file_type.lower()}" if not file_type.startswith('.') else file_type.lower()
    return file_ext in TextExtractor.SUPPORTED_EXTENSIONS